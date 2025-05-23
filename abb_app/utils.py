import os
import re
import regex
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import Counter
from docx.table import _Cell, Table
from docx.oxml.table import CT_Tbl
from typing import (
    TypedDict, Union, List, Dict, Set, Counter, Optional
)
from abb_app.model_integration.client import ModelClient
from thefuzz import process, fuzz


import logging

standard_logger = logging.getLogger('django')
standard_logger.error("STARTING TEXT EXTRACTION")

ABB_DICT_PATH = os.path.join(
    os.path.dirname(__file__),
    'data', 'abb_dict.csv'
)
SECTION_PATTERNS = [
    'ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОПРЕДЕЛЕНИЯ ТЕРМИНОВ', 'СПИСОК СОКРАЩЕНИЙ'
]
SKIP_SECTIONS = [
    "СПИСОК ЛИТЕРАТУРЫ", "Список использованной литературы",
    "Список использованных источников"
]
EXCLUDE_TERMS = {
    'ДИЗАЙН', 'ГЛАВНЫЙ', 'СПИСОК', 'ПРЯМОЙ', 'ПРИЕМ', 'ПРОТОКОЛ', 'ОТБОР',
    'КАЧЕСТВА', 'ПЕРИОД', 'ВЕДЕНИЕ', 'ЭТАП', 'ЭТИКА', 'СИНОПСИС', 'ЛИСТ',
    'ЦЕЛИ', 'РАБОТА', 'ИСТОРИЯ', 'ОЦЕНКА', 'СПОНСОР', 'ЗАДАЧИ', 'ДОСТУП',
    'КОНТРОЛЬ', 'ТЕРМИНОВ', 'ЗАПИСЕЙ', 'ГИПОТЕЗА', 'ДАННЫМИ', 'ДЕЙСТВИЕ',
    'ДАННЫМ/ДОКУМЕНТАЦИИ', 'ДЛЯ', 'ФОРМА', 'ВВЕДЕНИЕ', 'СВОЙСТВА', 'РЕЗЮМЕ',
    'ДАННЫХ', 'ЧЕЛОВЕКА', 'ОБЩЕСТВО', 'ЦЕНТР', 'АКТИВНЫХ', 'ВЕЩЕСТВ',
    'НАУЧНЫЙ', 'ОТЧЕТ', 'ОБЗОР', 'Каплана-Мейера', 'Стивенса-Джонсона',
    'Спрейг-Доули', 'Спрег-Доули', 'Мантеля-Хензеля', 'Нью-Йоркской',
    'Лонг-Эванс', 'ГмбХ', 'ТАБЛИЦ', 'РИСУНКОВ'
}

class Abbreviation(TypedDict):
    """Universal abbreviation structure"""
    abbreviation: str
    descriptions: List[str]  # All possible descriptions
    selected_description: Optional[str]  # User selected or entered description
    count: Optional[int]  # Number of occurrences in text
    contexts: Optional[List[str]]  # Context snippets
    correct_form: Optional[str]  # For mixed-character cases
    highlighted: Optional[List[Dict]]  # For display
    status: Optional[str]  # For tracking state
    is_ai_generated: bool # is the description was generated by model?

def load_abbreviation_dict() -> List[Abbreviation]:
    """Load abbreviations from CSV into list of Abbreviation"""
    abb_dict: Dict[str, List[str]] = {}
    
    if not os.path.exists(ABB_DICT_PATH):
        print(f"[WARNING] Abbreviation dictionary not found at: {ABB_DICT_PATH}")
        return []
        
    with open(ABB_DICT_PATH, 'r', encoding='utf-8') as f:
        next(f)
        for line in f:
            abb, desc = line.strip().split(',', 1)
            if abb in abb_dict:
                abb_dict[abb].append(desc)
            else:
                abb_dict[abb] = [desc]
    
    return [
        {
            'abbreviation': abb,
            'descriptions': descriptions
        }
        for abb, descriptions in abb_dict.items()
    ]

# -----------------------------------------------------------------------------
# Abbreviation table extraction
# -----------------------------------------------------------------------------

class AbbreviationTableExtractor:
    """Class for extracting abbreviation table from a Word document."""
    NS = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    def __init__(self, section_patterns: List[str] = SECTION_PATTERNS):
        self.section_patterns = section_patterns

    def get_abbreviation_table(self, doc: Document) -> List[Abbreviation]:
        """Extract abbreviations table from document"""
        table_element = self._extract_table_element(doc)
        if table_element is None:
            print(f"[INFO] No relevant abbreviation table found")
            return []
        
        print("[INFO] Found table with",
              f"{len(table_element.findall('.//w:tr', namespaces=self.NS))} rows"
              )
            
        return self._parse_table(table_element)
    
    def _extract_table_element(self, doc: Document) -> Optional[CT_Tbl]:
        """
        Extract the first table following a section matching `section_patterns`.
        Returns the block containing the table, or None if not found.
        """
        found_section = False
        
        for block in doc.element.body:
            if block.tag.endswith('p'): # paragraph block
                para_text = ''.join(
                    node.text for node in block.findall(
                        './/w:t', namespaces=self.NS
                    ) if node.text
                ).strip()

                if any(pattern.casefold() in para_text.casefold()
                    for pattern in self.section_patterns):
                    
                    # Must NOT have a hyperlink (avoid ToC lines)
                    if block.find('.//w:hyperlink', namespaces=self.NS) is not None:
                        continue
                    
                    # Must NOT end with a digit (avoid missformated ToC)
                    if para_text.strip().endswith(tuple("0123456789")):
                        continue
                    
                    # Must have some heading indication (pStyle or outlineLvl)
                    para_style = block.find('.//w:pStyle', namespaces=self.NS)
                    outline_level = block.find('.//w:outlineLvl', namespaces=self.NS)
                    if para_style is not None or outline_level is not None:
                        found_section = True
                        continue

            if found_section and block.tag.endswith('tbl'):
                return block

        return None
        
    def _parse_table(self, table_element: CT_Tbl) -> List[Abbreviation]:
        """Parse table into list of abbreviation entries"""
        abb_entries: Dict[str, List[str]] = {}
        rows = table_element.findall('.//w:tr', namespaces=self.NS)
        
        # TMP: debugging
        header_cells = rows[0].findall('.//w:tc', namespaces=self.NS)
        header_texts = [
            ''.join(
                t.text for t in cell.findall('.//w:t', namespaces=self.NS) if t.text
                ).strip()
            for cell in header_cells
        ]
        print(f"[INFO] Header: {header_texts}")
        
        # Get abbreviations and descriptions
        for idx, row in enumerate(rows):
            cell_values = [
                ''.join(
                    t.text for t in cell.findall('.//w:t', namespaces=self.NS) if t.text
                ).strip()
                for cell in row.findall('.//w:tc', namespaces=self.NS)
            ]
            if idx == 0 and set(cell_values[:2]) == {"Аббревиатура", "Расшифровка"}:
                continue

            if len(cell_values) == 2:
                abb, description = cell_values
                if abb in abb_entries:
                    if description not in abb_entries[abb]:
                        abb_entries[abb].append(description)
                else:
                    abb_entries[abb] = [description]
            else:
                print(f"[WARNING] Unexpected row structure: {cell_values}")

        return [
            {
                'abbreviation': abb,
                'descriptions': descriptions
            }
            for abb, descriptions in abb_entries.items()
        ]

# -----------------------------------------------------------------------------
# Text, abbreviation, and context extraction
# -----------------------------------------------------------------------------

class TextProcessor:
    """
    Class for extracting a relevant text, abbreviations and their contexts from
    a Word document.
    """
    def __init__(
            self,
            skip_sections: List[str] = SKIP_SECTIONS,
            exclude_terms: Set[str] = EXCLUDE_TERMS
        ):
        self.skip_sections = {
                section.upper() for section in skip_sections
            }
        self.exclude_terms = exclude_terms
        self.roman_pattern = re.compile(
            r'^(?:[IVXLCDM]+(?:-[IVXLCDM]+)?)[A-Za-zА-Яа-яёЁ]*$',
            re.IGNORECASE
        )

    def extract_relevant_text(self, doc: Document) -> str:
        """
        Extracts text from the document, excluding `skip_sections`.
        The exclusion starts at the section title in bold or heading style
        and resumes at the next bold or heading section.
        """
        paragraphs = []
        skip = False
        
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue
                
            try:
                style_name = None
                if hasattr(para, 'style') and para.style:
                    try:
                        style_name = para.style.name
                    except:
                        style_name = None
                        
                is_heading = (
                    style_name and 
                    (style_name.startswith('Heading') or 'Заголовок' in style_name)
                )
                
            except Exception as e:
                standard_logger.error(f"Error in para {i}: {str(e)}")
                is_heading = False
            
            is_bold = False
            if not is_heading:
                for run in para.runs:
                    if run.text.strip() and run.bold:
                        is_bold = True
                        break

            if (is_bold or is_heading):
                para_text_upper = para_text.upper()
                if any(section in para_text_upper for section in self.skip_sections):
                    skip = True
                elif skip:
                    skip = False

            if not skip:
                paragraphs.append(para_text)

        standard_logger.error("EXTRACTION COMPLETED")
        return ' '.join(paragraphs)

    def extract_abbreviations(self, text: str) -> Counter[str]:
        """
        Extracts uppercase and mixed-case abbreviations from the text.
        Excludes pure Roman numerals, `exclude_terms`, words in quotes,
        and words that are 9 or more characters long and contain only letters.
        """
        doc_abbs: Counter[str] = Counter()

        # Remove quoted words
        text_no_quotes = re.compile(r'«\S+?»|\"[^\"]+\"').sub('', text)
        words = text_no_quotes.split()

        # Find words with at least 2 uppercase (Latin or Cyrillic) letters
        matches = [
            word for word in words
            if re.search(r'[A-ZА-ЯЁ].*[A-ZА-ЯЁ]', word)
        ]
        
        for match in matches:
            clean_match = self._clean_abbreviation(match)

            # Exclude pure Roman numerals and specific excluded terms,
            # and overly long abbreviations
            if (not self.roman_pattern.match(clean_match)
                and clean_match not in self.exclude_terms
                and not (len(clean_match) > 8 and clean_match.isalpha())
                ):
                doc_abbs[clean_match] += 1
        print(f"Found {len(doc_abbs)} abbreviations")
        return doc_abbs

    def _clean_abbreviation(self, match: str) -> str:
        """Helper method to clean and format abbreviation matches."""
        clean_match = match.strip(':;,.»«][')

        # Remove '(' and ')' if unmatched, e.g. 'IgG)' but not in 'IgG(1)'
        if clean_match.startswith('('):
            clean_match = clean_match[1:]
        if clean_match.endswith(')') and clean_match.count('(') == 0:
            clean_match = re.sub(r'\)+$', '', clean_match)
        
        return clean_match.strip('»«][')

    def find_abbreviation_context(
            self,
            text: str,
            abbreviation: str,
            window: int = 50,
            max_contexts: int = 1000
        ) -> Union[str, List[str]]:
        """
        Finds and returns snippets of text around occurrences of the abbreviation.
        Limits the number of contexts returned to `max_contexts`.
        """
        contexts: Set[str] = set()
        matches = re.finditer(
            rf'(?<!\w){re.escape(abbreviation)}(?!\w)', text
        )
        for match in matches:
            start = max(0, match.start() - window)
            end = min(len(text), match.end() + window)
            snippet = "..." + text[start:end].strip() + "..."
            if max_contexts == 1:
                return [snippet]
            contexts.add(snippet)
            
        return list(contexts)

# -----------------------------------------------------------------------------
# Preparation of abbreviations
# -----------------------------------------------------------------------------

def process_abbreviations(
        doc: Document,
        abb_dict: List[Abbreviation]
    ) -> List[Abbreviation]:
    """Process abbreviations found in document"""
    text_processor = TextProcessor()
    validator = CharacterValidator()
    
    # Get abbreviations from document text
    text = text_processor.extract_relevant_text(doc)
    raw_abbs = text_processor.extract_abbreviations(text)    
    processed_abbs: List[Abbreviation] = []
    
    print("Starting processing abbreviations")

    for abb, count in raw_abbs.items():
        contexts = text_processor.find_abbreviation_context(text, abb)
        dict_entry = next(
            (entry for entry in abb_dict if entry['abbreviation'] == abb), None
        )
        
        descriptions = dict_entry['descriptions'] if dict_entry else []
        is_ai_generated = False        
            
        processed_abb: Abbreviation = {
            'abbreviation': abb,
            'descriptions': descriptions,
            'selected_description': None,  # Will be set by user
            'count': count,
            'contexts': contexts,
            'correct_form': None,
            'highlighted': None,
            'status': None,
            'is_ai_generated': is_ai_generated
        }
            
        # Validate and update if it's 9 or less characters long
        if len(abb) <= 15:
            try:
                val_result = validator.validate_abbreviation(abb, abb_dict)
                if val_result:
                    val_descriptions = val_result.get('descriptions', [])
                    processed_abb.update({
                        'correct_form': val_result.get('correct_form'),
                        'highlighted': val_result.get('highlighted'),
                        'descriptions': (
                            val_descriptions if val_descriptions 
                            else processed_abb['descriptions']
                        )
                    })
                    print(f"Processed validated abbreviation: {processed_abb}\n\n")
            except ValueError as e:
                print(f"[ERROR] Failed to validate abbreviation '{abb}': {e}")
            
        processed_abbs.append(processed_abb)
    
    return processed_abbs

# -----------------------------------------------------------------------------
# Formatting functions
# -----------------------------------------------------------------------------

class AbbreviationFormatter:
    """Class for formatting and cleaning abbreviation entries."""
    
    GREEK_TO_LATIN = {
        'α': 'A', 'β': 'B', 'γ': 'G', 'δ': 'D',
        'ε': 'E', 'ζ': 'Z', 'η': 'H', 'θ': 'TH',
        'ι': 'I', 'κ': 'K', 'λ': 'L', 'μ': 'M',
        'ν': 'N', 'ξ': 'X', 'ο': 'O', 'π': 'P',
        'ρ': 'R', 'σ': 'S', 'τ': 'T', 'υ': 'U',
        'φ': 'PH', 'χ': 'CH', 'ψ': 'PS', 'ω': 'O'
    }

    def format_description(self, entry: Dict[str, str]  ) -> str:
        """
        Format description by capitalizing words that correspond to abbreviation
        letters.
        """
        abbreviation: str = entry['abbreviation']
        description: str = entry['description']
        
        # Split description into English and Russian parts
        parts = description.split('(', 1)
        english_part = parts[0].strip().lower()
        russian_part = f"({parts[1]}" if len(parts) > 1 else ''
        
        # Convert abbreviation to uppercase Latin letters
        latin_abbr = ''.join(
            self.GREEK_TO_LATIN.get(char, char) for char in abbreviation
        ).upper()
        abbr_letters = ''.join(re.findall(r'[A-Z]', latin_abbr))
        
        english_part_capitalized = self._capitalize_by_abbreviation(
            english_part, abbr_letters
        )
        return f"{english_part_capitalized} {russian_part}".strip()

    def clean_and_sort_abbreviations(
            self, abbreviations: List[Dict[str, str]]
        ) -> List[Dict[str, str]]:
        """
        Clean and sort abbreviations:
        - Strips whitespace
        - Formats descriptions for abbreviations with English letters
        - Capitalizes first letters after digits
        - Removes duplicates
        - Sorts by abbreviation and description
        """
        # Create a copy to avoid modifying the original
        cleaned: List[Dict[str, str]] = []
        seen = set()  # For duplicate detection
        
        for entry in abbreviations:
            # Strip whitespace
            abb = entry['abbreviation'].strip()
            desc = entry['description'].strip()
            
            # Format if contains English letters
            if re.search(r'[A-Za-z]', abb):
                desc = self.format_description(
                    {'abbreviation': abb, 'description': desc}
                )
            
            # Capitalize after digits
            desc = self._capitalize_after_digits(desc)
            
            # Create unique key for deduplication
            unique_key = (abb, desc)
            if unique_key not in seen:
                seen.add(unique_key)
                cleaned.append({
                    'abbreviation': abb,
                    'description': desc
                })
        
        # Sort by abbreviation and description
        return sorted(cleaned, key=lambda x: (x['abbreviation'], x['description']))

    def _capitalize_by_abbreviation(
            self, text: str, abbr_letters: str
        ) -> str:
        """Capitalize words in text based on abbreviation letters."""
        abbr_index = 0  # Position in the abbreviation
        text_pos = 0  # Position in the text
        text_chars = list(text)

        while abbr_index < len(abbr_letters) and text_pos < len(text_chars):
            current_char = text_chars[text_pos]
            if (current_char.lower() == abbr_letters[abbr_index].lower()
                and (text_pos == 0 or not text_chars[text_pos - 1].isalpha())):
                text_chars[text_pos] = current_char.upper()
                abbr_index += 1
            text_pos += 1
            
        return ''.join(text_chars)

    def _capitalize_after_digits(self, text: str) -> str:
        """Capitalize the first letter following any leading digits."""
        return re.sub(
            r'^(\d*)([a-zA-ZА-Яа-яЁё])',
            lambda m: m.group(1) + m.group(2).upper(),
            text
        )

# -----------------------------------------------------------------------------
# Mistyped Abbreviation Checking
# -----------------------------------------------------------------------------

class CharacterValidator:
    def __init__(self):
        # Map for character-by-character conversion
        self.cyr2lat = {
            'А': 'A', 'В': 'B', 'С': 'C', 'Е': 'E',
            'Н': 'H', 'К': 'K', 'М': 'M', 'О': 'O',
            'Р': 'P', 'Т': 'T', 'У': 'Y', 'Х': 'X'
        }
        # Add lowercase mappings
        self.cyr2lat.update({k.lower(): v.lower() 
                            for k, v in self.cyr2lat.items()})
        # Create reverse mapping
        self.lat2cyr = {v: k for k, v in self.cyr2lat.items()}

    def validate_abbreviation(
            self, 
            abb: str, 
            abb_dict: List[Abbreviation]
        ) -> dict:
        """
        Validates an abbreviation for mixed characters.
        Checks for existing forms in the dictionary.
        Returns a dict with validation info or empty dict.

        Decision Tree (important returns are shown):
        Abbreviation
        └─ has_cyr_chars OR has_lat_chars
            ├─ Generate forms and search the dictionary
            │    ├─ Match found (does not matter mixed or not)
            │    │    ├─ correct_form = matched_form
            │    │    ├─ descriptions = matched_description
            │    │    └─ highlighted = highlighted_text for tooltip
            │    └─ No match found
            │         ├─ is mixed (has_cyr_chars AND has_lat_chars)
            │         │     └─ highlighted = highlighted_text for moderation
            └─ Does not contain both types of characters or not mixed
                └─ no validation issues found, return empty dict

        Where:
        - `cyr_chars` and `lat_chars` refer to similar-looking
        Cyrillic and Latin characters.
        """
        has_cyr_chars = any(char in self.cyr2lat for char in abb)
        has_lat_chars = any(char in self.lat2cyr for char in abb)
        
        if not (has_cyr_chars or has_lat_chars):
            return {}
    
        # Generate all possible forms and search dictionary
        possible_forms = self._generate_all_mixed_forms(abb)
        matched_entries = [
            entry for entry in abb_dict 
            if entry['abbreviation'] in possible_forms
        ]
    
        if matched_entries:
            # Check for multiple matches
            unique_forms = set(entry['abbreviation'] for entry in matched_entries)
            if len(unique_forms) > 1:
                raise ValueError(
                    "[ERROR] Mixed-character abbreviations in the dictionary:"
                    f"\n{matched_entries}"
                )
            
            matched_entry = matched_entries[0]
            return {
                "correct_form": matched_entry['abbreviation'],
                "descriptions": matched_entry['descriptions'],
                "highlighted": self._highlight_mismatch_characters(
                    abb, matched_entry['abbreviation']
                )
            }

        # Handle case when no dictionary match found but chars are mixed
        if has_cyr_chars and has_lat_chars:
            return {
                "correct_form": None,
                "descriptions": [],
                "highlighted": self._highlight_mixed_characters(abb)
            }

        return {}

    def _generate_all_mixed_forms(self, abb: str) -> set:
        """Generate all possible character combinations"""
        results = set()
        
        # Add full conversions
        results.add("".join(self.lat2cyr.get(ch, ch) for ch in abb))
        results.add("".join(self.cyr2lat.get(ch, ch) for ch in abb))
        
        # Generate partial conversions
        def backtrack(i: int, current: list):
            if i == len(abb):
                results.add("".join(current))
                return

            ch = abb[i]
            # Original character
            current.append(ch)
            backtrack(i + 1, current)
            current.pop()

            # Convert if possible
            if ch in self.cyr2lat:
                current.append(self.cyr2lat[ch])
                backtrack(i + 1, current)
                current.pop()
            if ch in self.lat2cyr:
                current.append(self.lat2cyr[ch])
                backtrack(i + 1, current)
                current.pop()

        backtrack(0, [])
        return results - {abb}  # Exclude original form

    def _highlight_mismatch_characters(
            self, user_abb: str, dict_abb: str
            ) -> str:
        """
        Compare each character and return a list of dictionaries
        with mismatch information for template rendering.
        """
        highlighted = []
        for ch_user, ch_dict in zip(user_abb, dict_abb):
            if ch_user != ch_dict:
                mismatch_type = (
                    "кириллическая" if ch_user in self.cyr2lat else "латинская"
                )
                correct_type = (
                    "латинская" if ch_dict in self.lat2cyr else "кириллическая"
                )
                tooltip_text = (
                    f"{ch_user} - {mismatch_type}, "
                    f"в словаре {ch_dict} - {correct_type}"
                )
                highlighted.append({
                    "char": ch_user,
                    "tooltip": tooltip_text,
                    "mismatch": True
                })
            else:
                highlighted.append({
                    "char": ch_user,
                    "mismatch": False
                })
        return highlighted
    
    def _highlight_mixed_characters(self, abb: str) -> str:
        """Marks each character with its script type"""
        highlighted = []
        for ch in abb:
            if ch in self.cyr2lat:
                highlighted.append(f"{ch}(Cyr)")
            elif ch in self.lat2cyr:
                highlighted.append(f"{ch}(Lat)")
            else:
                highlighted.append(ch)
        return "".join(highlighted)

# -----------------------------------------------------------------------------
# Abbreviation comparison
# -----------------------------------------------------------------------------

def compare_abbreviations(
        old_abbs: List[Abbreviation],
        new_abbs: Union[List[Dict[str, str]], List[Abbreviation]],
    ) -> Dict[str, List[Abbreviation]]:
    """
    Compare new and old abbreviation tables.
    Returns dictionary with:
        - 'missing_abbs': abbreviations present in old but not in new
        - 'new_found': abbreviations present in new but not in old
    """
    results: Dict[str, List[Abbreviation]] = {}

    # Create sets of abbreviations for efficient comparison
    old_abb_set = {abb['abbreviation'] for abb in old_abbs}
    if new_abbs and isinstance(new_abbs[0], dict):
        if 'description' in new_abbs[0]:
            get_descriptions = lambda abb: [abb['description']]  # Wrap in list
        elif 'descriptions' in new_abbs[0]:
            get_descriptions = lambda abb: abb['descriptions']  # Use as is
        else:
            raise ValueError("Invalid format for new_abbs")

    new_abb_set = {abb['abbreviation'] for abb in new_abbs}

    # Find abbreviations that are in old but not in new
    results['missing_abbs'] = [
        abb for abb in old_abbs 
        if abb['abbreviation'] not in new_abb_set
    ]
    # Wrap description in a list to match the structure of old_abbs
    results['new_found'] = [
        {
            'abbreviation': abb['abbreviation'],
            'descriptions': get_descriptions(abb)
        }
        for abb in new_abbs
        if abb['abbreviation'] not in old_abb_set
    ]
    
    return results

# -----------------------------------------------------------------------------
# Output generation
# -----------------------------------------------------------------------------

class AbbreviationTableGenerator:
    """
    Class for generating formatted Word document tables with abbreviations.
    """

    def __init__(self):
        self.margins = {
            'top': 2.0,    # cm
            'bottom': 2.0, # cm
            'left': 3.0,   # cm
            'right': 1.5   # cm
        }
        self.first_column_width = 3.7  # cm
        self.font_name = 'Times New Roman'
        self.font_size = 12  # pt

    def generate_document(self, table_entries: List[Abbreviation]) -> Document:
        """
        Generate a Word document with formatted abbreviation table.
        """
        doc = Document()

        # Set page margins
        for section in doc.sections:
            section.top_margin = Cm(self.margins['top'])
            section.bottom_margin = Cm(self.margins['bottom'])
            section.left_margin = Cm(self.margins['left'])
            section.right_margin = Cm(self.margins['right'])

            # Calculate second column width
            total_width = (
                section.page_width 
                - section.left_margin 
                - section.right_margin
            )
            self.second_column_width = total_width - Cm(self.first_column_width)

        # Create and format table
        table = self._create_table(doc, table_entries)
        self._set_column_widths(table)

        return doc

    def _create_table(self, doc: Document, table_entries: List[Abbreviation]) -> Table:
        """Create and format table with header and data rows."""
        # Create table with header
        table = doc.add_table(rows=1, cols=2)
        header_cells = table.rows[0].cells
        
        # Set header text
        header_cells[0].text = 'Аббревиатура'
        header_cells[1].text = 'Расшифровка'

        # Format header cells
        for cell in header_cells:
            self._format_cell(cell, bold=True)

        # Add and format data rows
        for entry in table_entries:
            row_cells = table.add_row().cells
            row_cells[0].text = entry['abbreviation']
            row_cells[1].text = entry['description']

            for cell in row_cells:
                self._format_cell(cell, bold=False)

        return table
    
    def _set_cell_border(self, cell: _Cell) -> None:
        """Set black borders on all four sides of the given cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        for edge in ("top", "bottom", "left", "right"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn('w:val'), 'single')    # single line
            element.set(qn('w:sz'), '4')          # line size
            element.set(qn('w:color'), '000000')  # black color
            tcPr.append(element)

    def _format_paragraph_spacing(self, cell: _Cell) -> None:
        """Set line spacing to single and remove spacing before/after paragraphs."""
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = Pt(self.font_size)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

    def _format_cell_text(self, cell: _Cell, bold: bool = False) -> None:
        """Format text in the given cell with specified font and style."""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
                run.font.bold = bold
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def _format_cell(self, cell: _Cell, bold: bool = False) -> None:
        """Apply all formatting to a cell."""
        self._set_cell_border(cell)
        self._format_paragraph_spacing(cell)
        self._format_cell_text(cell, bold)

    def _set_column_widths(self, table: Table) -> None:
        """Set fixed widths for table columns."""
        for row in table.rows:
            row.cells[0].width = Cm(self.first_column_width)
            row.cells[1].width = self.second_column_width


# -----------------------------------------------------------------------------
# Alphabet detection for cleaning the abbreviation dictionary
# -----------------------------------------------------------------------------

def detect_string_alphabet(text):
    """Detect if string contains Russian, Latin or mixed characters.
    Returns: 'russian', 'latin', or 'mixed'
    """
    has_russian = bool(regex.search(r'\p{Script=Cyrillic}', text))
    has_latin   = bool(regex.search(r'\p{Script=Latin}', text))
    
    if has_russian and has_latin:
        return 'mixed'
    elif has_russian:
        return 'russian'
    elif has_latin:
        return 'latin'
    return 'other'

def split_by_language(text):
    """Split text into Russian and Latin parts while preserving compound terms.
    Returns: (russian_text, latin_text)
    """
    russian_parts = []
    latin_parts = []

    pattern = r'[\p{Script=Cyrillic}\p{Script=Latin}\p{N}_-]+[.,;:!?]*|\S'
    words = regex.findall(pattern, text)
    # print(words)

    for word in words:
        lang = detect_string_alphabet(word)
        if lang == 'russian':
            russian_parts.append(word)
        elif lang == 'latin':
            latin_parts.append(word)

    russian_text = ' '.join(russian_parts)
    latin_text = ' '.join(latin_parts)

    # Remove trailing punctuation
    russian_text = regex.sub(r'[.,;:!?]+$', '', russian_text)
    latin_text = regex.sub(r'[.,;:!?]+$', '', latin_text)

    return russian_text, latin_text