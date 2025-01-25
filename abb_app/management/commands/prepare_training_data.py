import os
import csv
import random
from django.core.management.base import BaseCommand
from docx import Document
from docx.shared import Pt, RGBColor, Cm

class Command(BaseCommand):
    help = 'Prepare training data by splitting into train and test sets'

    def add_arguments(self, parser):
        parser.add_argument(
            '--input-file',
            type=str,
            default=os.path.join('abb_app', 'data', 'abb_dict_with_contexts.csv'),
            help='Input CSV file with abbreviations and contexts'
        )
        parser.add_argument(
            '--train-file',
            type=str,
            default=os.path.join('abb_app', 'data', 'train_data.csv'),
            help='Output file for training data'
        )
        parser.add_argument(
            '--test-file',
            type=str,
            default=os.path.join('abb_app', 'data', 'test_data.docx'),
            help='Output Word document for test data'
        )
        parser.add_argument(
            '--test-size',
            type=float,
            default=0.2,
            help='Proportion of data to use for testing (0.0-1.0)'
        )
        parser.add_argument(
            '--overlap-size',
            type=int,
            default=5,
            help='Number of training examples to include in test set'
        )
        parser.add_argument(
            '--seed',
            type=int,
            default=42,
            help='Random seed for reproducibility'
        )

    def _create_test_document(self, test_data, overlap_examples):
        """Create a formatted Word document with test data"""
        doc = Document()
        
        # Set page margins
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(1.5)

        # Add title
        doc.add_heading('Example set for testing LLM-based description generation', 0)

        # Add overlap examples section if any
        if overlap_examples:
            doc.add_paragraph('Examples from training set:')
            doc.add_paragraph()
            
            for item in overlap_examples:
                para = doc.add_paragraph()
                para.add_run(f"{item['abbreviation']} - {item['description']}")
                if item['contexts']:
                    context = item['contexts'].split(' ... ')[0]
                    doc.add_paragraph(f"Example usage: {context}")
                doc.add_paragraph()

        # Add new test examples
        doc.add_paragraph('New examples for testing:')
        doc.add_paragraph()
        
        for item in test_data:
            para = doc.add_paragraph()
            para.add_run(f"{item['abbreviation']} - {item['description']}")
            if item['contexts']:
                context = item['contexts'].split(' ... ')[0]
                doc.add_paragraph(f"Example usage: {context}")
            doc.add_paragraph()

        return doc

    def handle(self, *args, **options):
        input_file = options['input_file']
        train_file = options['train_file']
        test_file = options['test_file']
        test_size = options['test_size']
        overlap_size = options['overlap_size']
        random.seed(options['seed'])

        if not os.path.exists(input_file):
            self.stderr.write(f"Input file not found: {input_file}")
            return

        # Read all data
        all_data = []
        with open(input_file, 'r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            header = next(reader)
            for row in reader:
                if len(row) == 3:
                    all_data.append({
                        'abbreviation': row[0],
                        'description': row[1],
                        'contexts': row[2]
                    })

        total_examples = len(all_data)
        if total_examples == 0:
            self.stderr.write("No valid data found in input file")
            return

        # Shuffle and split data
        random.shuffle(all_data)
        test_count = int(total_examples * test_size)
        
        # Select test data
        test_data = all_data[:test_count]
        train_data = all_data[test_count:]

        # Select overlap examples from training data
        overlap_examples = []
        if overlap_size > 0 and len(train_data) >= overlap_size:
            overlap_examples = random.sample(train_data, overlap_size)

        # Save train data
        os.makedirs(os.path.dirname(train_file), exist_ok=True)
        with open(train_file, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.writer(f, quoting=csv.QUOTE_ALL)
            writer.writerow(header)
            for item in train_data:
                writer.writerow([
                    item['abbreviation'],
                    item['description'],
                    item['contexts']
                ])

        # Create and save test document
        doc = self._create_test_document(test_data, overlap_examples)
        os.makedirs(os.path.dirname(test_file), exist_ok=True)
        doc.save(test_file)

        self.stdout.write(self.style.SUCCESS(
            f"\nData split complete:\n"
            f"- Training examples: {len(train_data)}\n"
            f"- Test examples: {len(test_data)} "
            f"(including {len(overlap_examples)} from training)\n"
            f"Files saved:\n"
            f"- Training data: {train_file}\n"
            f"- Test data: {test_file}"
        )) 