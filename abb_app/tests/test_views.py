import io
import json
from pathlib import Path
from tempfile import TemporaryDirectory

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from django.urls import reverse
from docx import Document

from abb_app.models import AbbreviationEntry
from abb_app.services.sessions import (
    CARD_STATE_SESSION_KEY,
    DEMO_FILENAME,
    PROCESSED_FILE_SESSION_KEY,
    SESSION_FILE_KEY,
    TABLE_CHECK_SESSION_KEY,
    WORKFLOW_STATE_SESSION_KEY,
)


class ProcessingViewTests(TestCase):
    def test_document_is_processed_and_stored_in_session(self):
        AbbreviationEntry.objects.create(
            abbreviation='T4',
            description='thyroxine',
            status='approved',
        )
        with TemporaryDirectory() as media_root:
            session_id = 'test-session'
            path = Path(media_root) / f'{session_id}.docx'
            doc = Document()
            doc.add_paragraph('У пациента определяли уровень T4.')
            doc.save(path)

            session = self.client.session
            session[SESSION_FILE_KEY] = path.name
            session.save()
            with self.settings(MEDIA_ROOT=media_root):
                response = self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )

            self.assertEqual(response.status_code, 200)
            self.assertTemplateUsed(response, 'content.html')
            doc_abbs = self.client.session['doc_abbs']
            self.assertEqual(len(doc_abbs), 1)
            self.assertEqual(doc_abbs[0]['abbreviation'], 'T4')
            self.assertEqual(doc_abbs[0]['descriptions'], ['thyroxine'])
            self.assertEqual(
                self.client.session[PROCESSED_FILE_SESSION_KEY],
                path.name,
            )

    def test_refresh_preserves_review_card_and_comparison_state(self):
        AbbreviationEntry.objects.create(
            abbreviation='T4',
            description='thyroxine',
            status='approved',
        )
        with TemporaryDirectory() as media_root:
            session_id = 'review-state'
            path = Path(media_root) / f'{session_id}.docx'
            doc = Document()
            doc.add_paragraph('У пациента определяли уровень T4.')
            doc.save(path)
            session = self.client.session
            session[SESSION_FILE_KEY] = path.name
            session.save()

            with self.settings(MEDIA_ROOT=media_root):
                self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )
                session = self.client.session
                session['initial_abbs'] = [{
                    'abbreviation': 'T4',
                    'descriptions': ['thyroxine'],
                }]
                session[TABLE_CHECK_SESSION_KEY] = True
                session.save()
                before_review = self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )
                update = self.client.post(
                    reverse('update_abbreviation'),
                    data=json.dumps({
                        'abbreviation': 'T4',
                        'description': 'thyroxine',
                        'action': 'add',
                    }),
                    content_type='application/json',
                )
                refreshed = self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )

            self.assertTrue(
                before_review.context['comparison_context'][
                    'waiting_for_review'
                ]
            )
            self.assertEqual(update.status_code, 200)
            entry = self.client.session['doc_abbs'][0]
            self.assertEqual(entry['selected_description'], 'thyroxine')
            self.assertTrue(entry['reviewed'])
            self.assertIn(
                'T4',
                self.client.session[CARD_STATE_SESSION_KEY],
            )
            self.assertIn(
                'T4',
                refreshed.context['collapsed_abbreviations'],
            )
            self.assertNotIn(
                'waiting_for_review',
                refreshed.context['comparison_context'],
            )

    def test_refresh_preserves_skipped_and_manual_card_state(self):
        with TemporaryDirectory() as media_root:
            session_id = 'card-state'
            path = Path(media_root) / f'{session_id}.docx'
            Document().save(path)
            session = self.client.session
            session[SESSION_FILE_KEY] = path.name
            session[PROCESSED_FILE_SESSION_KEY] = path.name
            session['doc_abbs'] = [
                {
                    'abbreviation': 'T4',
                    'descriptions': ['thyroxine'],
                    'selected_description': None,
                    'reviewed': False,
                    'occurrence_count': 2,
                },
                {
                    'abbreviation': 'ABC',
                    'descriptions': [],
                    'selected_description': None,
                    'reviewed': False,
                    'occurrence_count': 2,
                },
            ]
            session['initial_abbs'] = []
            session.save()

            with self.settings(MEDIA_ROOT=media_root):
                skipped = self.client.post(
                    reverse('update_abbreviation'),
                    data=json.dumps({
                        'abbreviation': 'T4',
                        'action': 'skip',
                    }),
                    content_type='application/json',
                )
                collapsed = self.client.post(
                    reverse('update_card_state'),
                    data=json.dumps({
                        'abbreviation': 'ABC',
                        'collapsed': True,
                    }),
                    content_type='application/json',
                )
                refreshed = self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )

            self.assertEqual(skipped.status_code, 200)
            self.assertEqual(collapsed.status_code, 200)
            entries = {
                item['abbreviation']: item
                for item in self.client.session['doc_abbs']
            }
            self.assertTrue(entries['T4']['reviewed'])
            self.assertIsNone(entries['T4']['selected_description'])
            self.assertEqual(
                refreshed.context['collapsed_abbreviations'],
                {'T4', 'ABC'},
            )

    def test_refresh_preserves_table_choice_and_workflow_state(self):
        with TemporaryDirectory() as media_root:
            session_id = 'workflow-state'
            path = Path(media_root) / f'{session_id}.docx'
            Document().save(path)
            session = self.client.session
            session[SESSION_FILE_KEY] = path.name
            session[PROCESSED_FILE_SESSION_KEY] = path.name
            session['doc_abbs'] = []
            session['initial_abbs'] = [{
                'abbreviation': 'T4',
                'description': 'thyroxine',
            }]
            session.save()

            with self.settings(MEDIA_ROOT=media_root):
                self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )
                disabled = self.client.post(
                    reverse('update_table_check'),
                    data=json.dumps({'enabled': False}),
                    content_type='application/json',
                )
                disabled_refresh = self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )
                enabled = self.client.post(
                    reverse('update_table_check'),
                    data=json.dumps({'enabled': True}),
                    content_type='application/json',
                )
                sizes = {
                    'comparison-block': 240,
                    'table-preview-tool': 320,
                }
                for tool_id, height in sizes.items():
                    response = self.client.post(
                        reverse('update_workflow_state'),
                        data=json.dumps({
                            'tool_id': tool_id,
                            'open': True,
                            'height': height,
                        }),
                        content_type='application/json',
                    )
                    self.assertEqual(response.status_code, 200)
                refreshed = self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )

            self.assertEqual(disabled.status_code, 200)
            self.assertFalse(disabled_refresh.context['table_check_enabled'])
            self.assertEqual(enabled.status_code, 200)
            self.assertTrue(refreshed.context['table_check_enabled'])
            self.assertTrue(refreshed.context['comparison_open'])
            self.assertEqual(refreshed.context['comparison_height'], 240)
            self.assertTrue(refreshed.context['table_preview_open'])
            self.assertEqual(refreshed.context['table_preview_height'], 320)

    def test_new_upload_resets_processing_state(self):
        buffer = io.BytesIO()
        Document().save(buffer)
        upload = SimpleUploadedFile('new.docx', buffer.getvalue())
        session = self.client.session
        session[TABLE_CHECK_SESSION_KEY] = True
        session[CARD_STATE_SESSION_KEY] = ['OLD']
        session[WORKFLOW_STATE_SESSION_KEY] = {
            'comparison-block': {'open': True, 'height': 240},
            'table-preview-tool': {'open': True, 'height': 320},
        }
        session[PROCESSED_FILE_SESSION_KEY] = 'old.docx'
        session['doc_abbs'] = [{'abbreviation': 'OLD'}]
        session['initial_abbs'] = [{'abbreviation': 'OLD'}]
        session.save()
        with TemporaryDirectory() as media_root:
            with self.settings(MEDIA_ROOT=media_root):
                response = self.client.post(
                    reverse('upload_file'),
                    {'uploaded_file': upload},
                )

        self.assertEqual(response.status_code, 302)
        session = self.client.session
        for key in (
            TABLE_CHECK_SESSION_KEY,
            WORKFLOW_STATE_SESSION_KEY,
            CARD_STATE_SESSION_KEY,
            PROCESSED_FILE_SESSION_KEY,
            'doc_abbs',
            'initial_abbs',
        ):
            with self.subTest(key=key):
                self.assertNotIn(key, session)

    def test_demo_refresh_starts_from_clean_processing_state(self):
        AbbreviationEntry.objects.create(
            abbreviation='T4',
            description='thyroxine',
            status='approved',
        )
        with TemporaryDirectory() as media_root:
            path = Path(media_root) / DEMO_FILENAME
            doc = Document()
            doc.add_paragraph('T4')
            doc.save(path)
            session = self.client.session
            session[SESSION_FILE_KEY] = DEMO_FILENAME
            session[PROCESSED_FILE_SESSION_KEY] = DEMO_FILENAME
            session['doc_abbs'] = [{
                'abbreviation': 'OLD',
                'selected_description': 'old',
                'reviewed': True,
            }]
            session['initial_abbs'] = [{'abbreviation': 'OLD'}]
            session[TABLE_CHECK_SESSION_KEY] = False
            session[CARD_STATE_SESSION_KEY] = ['OLD']
            session[WORKFLOW_STATE_SESSION_KEY] = {
                'comparison-block': {'open': True, 'height': 240},
            }
            session.save()

            with self.settings(MEDIA_ROOT=media_root):
                response = self.client.get(
                    reverse('process_file_with_session', args=['test_drive'])
                )

            self.assertEqual(response.status_code, 200)
            self.assertTrue(response.context['is_demo'])
            self.assertTrue(response.context['table_check_enabled'])
            self.assertEqual(response.context['collapsed_abbreviations'], set())
            self.assertFalse(response.context['comparison_open'])
            self.assertNotIn(TABLE_CHECK_SESSION_KEY, self.client.session)
            self.assertNotIn(WORKFLOW_STATE_SESSION_KEY, self.client.session)
            self.assertEqual(
                [
                    entry['abbreviation']
                    for entry in self.client.session['doc_abbs']
                ],
                ['T4'],
            )
            self.assertIsNone(
                self.client.session['doc_abbs'][0]['selected_description']
            )

    def test_bibliography_review_precedes_processing_and_table_check(self):
        with TemporaryDirectory() as media_root:
            session_id = 'bibliography'
            path = Path(media_root) / f'{session_id}.docx'
            document = Document()
            document.add_heading('СПИСОК СОКРАЩЕНИЙ', level=1)
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'Аббревиатура'
            table.cell(0, 1).text = 'Расшифровка'
            table.cell(1, 0).text = 'ABC'
            table.cell(1, 1).text = 'alpha beta complex'
            document.add_paragraph('Основной текст ABC ABC.')
            document.add_paragraph('12: ЛИТЕРАТУРА')
            document.add_paragraph(
                'Akerman S, Goadsby PJ. WHO study. 2005;146(1):7-14. '
                'doi:10.1/first'
            )
            document.add_heading('Вторая часть', level=1)
            document.add_paragraph('Основной текст DEF DEF.')
            document.add_heading('Публикации по второй части', level=1)
            document.add_paragraph(
                'Smith A, Jones BC. XYZ study. 2021;12(1):123-130. '
                'doi:10.1/second'
            )
            document.add_paragraph(
                'Brown D, White EF. XYZ study. 2022;13(2):131-140. '
                'doi:10.1/third'
            )
            document.add_heading('Заключение', level=1)
            document.add_paragraph('Итоговый текст END END.')
            document.save(path)
            session = self.client.session
            session[SESSION_FILE_KEY] = path.name
            session.save()

            process_url = reverse(
                'process_file_with_session',
                args=[session_id],
            )
            with self.settings(MEDIA_ROOT=media_root):
                review = self.client.get(process_url)
                sections = review.context['bibliography_sections']
                first_section_id = sections[0].section_id
                self.assertEqual(review.status_code, 200)
                self.assertTemplateUsed(
                    review,
                    'bibliography_review.html',
                )
                self.assertEqual(
                    [section.title for section in sections],
                    [
                        '12: ЛИТЕРАТУРА',
                        'Публикации по второй части',
                    ],
                )
                self.assertNotIn('doc_abbs', self.client.session)
                processed = self.client.post(
                    process_url,
                    data={
                        'include_bibliography_sections': first_section_id,
                    },
                    follow=True,
                )

            self.assertEqual(processed.status_code, 200)
            self.assertEqual(
                processed.redirect_chain,
                [(process_url, 302)],
            )
            self.assertTemplateUsed(processed, 'content.html')
            self.assertTrue(processed.context['show_table_check_dialog'])
            entries = {
                entry['abbreviation']: entry
                for entry in self.client.session['doc_abbs']
            }
            self.assertIn('WHO', entries)
            self.assertNotIn('XYZ', entries)
            self.assertEqual(
                self.client.session[PROCESSED_FILE_SESSION_KEY],
                path.name,
            )

    def test_frequency_groups_and_bulk_toggle_singletons(self):
        with TemporaryDirectory() as media_root:
            session_id = 'singletons'
            path = Path(media_root) / f'{session_id}.docx'
            document = Document()
            document.add_paragraph('ABC ABC АБВ СD СD')
            document.save(path)

            session = self.client.session
            session[SESSION_FILE_KEY] = path.name
            session.save()
            with self.settings(MEDIA_ROOT=media_root):
                response = self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )
                selected = self.client.post(
                    reverse('update_abbreviation'),
                    data=json.dumps({
                        'abbreviation': 'АБВ',
                        'description': 'тестовая расшифровка',
                        'action': 'add',
                    }),
                    content_type='application/json',
                )
                removed = self.client.post(
                    reverse('update_single_occurrence_abbreviations'),
                    data=json.dumps({'action': 'remove'}),
                    content_type='application/json',
                )

            self.assertEqual(
                [
                    entry['abbreviation']
                    for entry in response.context['repeated_abbs']
                ],
                ['ABC', 'СD'],
            )
            self.assertEqual(
                [
                    entry['abbreviation']
                    for entry in response.context['single_occurrence_abbs']
                ],
                ['АБВ'],
            )
            self.assertEqual(selected.status_code, 200)
            self.assertEqual(removed.status_code, 204)
            removed_entry = next(
                entry for entry in self.client.session['doc_abbs']
                if entry['abbreviation'] == 'АБВ'
            )
            self.assertTrue(removed_entry['reviewed'])
            self.assertIsNone(removed_entry['selected_description'])
            self.assertIn(
                'АБВ',
                self.client.session.get(CARD_STATE_SESSION_KEY, []),
            )

            with self.settings(MEDIA_ROOT=media_root):
                added = self.client.post(
                    reverse('update_single_occurrence_abbreviations'),
                    data=json.dumps({'action': 'add'}),
                    content_type='application/json',
                )

            self.assertEqual(added.status_code, 204)
            entries = {
                entry['abbreviation']: entry
                for entry in self.client.session['doc_abbs']
            }
            self.assertEqual(entries['ABC']['occurrence_count'], 2)
            self.assertEqual(entries['СD']['occurrence_count'], 2)
            self.assertEqual(entries['АБВ']['occurrence_count'], 1)
            self.assertFalse(entries['АБВ']['reviewed'])
            self.assertIsNone(entries['АБВ']['selected_description'])
            self.assertNotIn(
                'АБВ',
                self.client.session.get(CARD_STATE_SESSION_KEY, []),
            )

    def test_document_url_requires_matching_django_session(self):
        with TemporaryDirectory() as media_root:
            session_id = 'test-session'
            path = Path(media_root) / f'{session_id}.docx'
            Document().save(path)

            with self.settings(MEDIA_ROOT=media_root):
                response = self.client.get(
                    reverse('process_file_with_session', args=[session_id])
                )

            self.assertEqual(response.status_code, 302)
            self.assertEqual(response.url, reverse('upload_file'))
