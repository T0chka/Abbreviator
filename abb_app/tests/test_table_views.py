import io
import json
from pathlib import Path
from tempfile import TemporaryDirectory

from django.test import TestCase
from django.urls import reverse
from docx import Document

from abb_app.models import AbbreviationEntry
from abb_app.services.documents import process_document
from abb_app.services.extraction import CharacterValidator


CANONICAL = 'CTCAE'
MIXED = 'СTCAE'
DESCRIPTION = 'common terminology criteria for adverse events'


def mixed_card():
    validation = CharacterValidator().validate_abbreviation(
        MIXED,
        [{
            'abbreviation': CANONICAL,
            'descriptions': [DESCRIPTION],
        }],
    )
    return {
        'abbreviation': MIXED,
        'descriptions': [DESCRIPTION],
        'selected_description': DESCRIPTION,
        'reviewed': True,
        'correct_form': validation['correct_form'],
        'highlighted': validation['highlighted'],
    }


def canonical_card():
    return {
        'abbreviation': CANONICAL,
        'descriptions': [DESCRIPTION],
        'selected_description': DESCRIPTION,
        'reviewed': True,
        'correct_form': None,
        'highlighted': None,
    }


def initial_entry(abbreviation, highlighted=None):
    return {
        'abbreviation': abbreviation,
        'descriptions': [DESCRIPTION],
        'highlighted': highlighted,
    }


class InitialTableValidationTests(TestCase):
    def test_initial_table_uses_same_mixed_alphabet_validation_as_cards(self):
        AbbreviationEntry.objects.create(
            abbreviation=CANONICAL,
            description=DESCRIPTION,
            status='approved',
        )

        with TemporaryDirectory() as directory:
            path = Path(directory) / 'initial-table.docx'
            doc = Document()
            doc.add_heading('СПИСОК СОКРАЩЕНИЙ', level=1)
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = 'Аббревиатура'
            table.rows[0].cells[1].text = 'Расшифровка'
            table.rows[1].cells[0].text = MIXED
            table.rows[1].cells[1].text = DESCRIPTION
            doc.save(path)

            processed = process_document(str(path))

        initial = processed.initial_abbreviations[0]
        self.assertEqual(initial['abbreviation'], MIXED)
        self.assertTrue(
            any(part['mismatch'] for part in initial['highlighted'])
        )


class TableComparisonViewTests(TestCase):
    def set_session(self, doc_abbs, initial_abbs):
        session = self.client.session
        session['doc_abbs'] = doc_abbs
        session['initial_abbs'] = initial_abbs
        session.save()

    def post_comparison(self, use_correct_form, scope='all'):
        return self.client.post(
            reverse('update_difference_section'),
            data=json.dumps({
                'use_correct_form': use_correct_form,
                'scope': scope,
            }),
            content_type='application/json',
        )

    def test_comparison_ignores_scope_and_highlights_original_spelling(self):
        self.set_session(
            [mixed_card()],
            [initial_entry(CANONICAL)],
        )

        off_all = self.post_comparison(False, 'all')
        off_new = self.post_comparison(False, 'new')
        on_all = self.post_comparison(True, 'all')
        on_new = self.post_comparison(True, 'new')

        self.assertEqual(off_all.status_code, 200)
        self.assertEqual(off_all.content, off_new.content)
        self.assertContains(off_all, 'tooltip tooltip-right red')
        self.assertEqual(on_all.content, on_new.content)
        self.assertNotContains(on_all, 'tooltip tooltip-right red')

        validation = CharacterValidator().validate_abbreviation(
            MIXED,
            [{
                'abbreviation': CANONICAL,
                'descriptions': [DESCRIPTION],
            }],
        )
        self.set_session(
            [canonical_card()],
            [initial_entry(MIXED, validation['highlighted'])],
        )
        missing_original = self.post_comparison(True)

        self.assertEqual(missing_original.status_code, 200)
        self.assertContains(
            missing_original,
            'tooltip tooltip-right red',
        )


class PreviewExportParityTests(TestCase):
    def test_preview_and_export_use_the_same_rows(self):
        scenarios = [
            {
                'initial': [],
                'settings': {
                    'use_correct_form': False,
                    'scope': 'all',
                },
            },
            {
                'initial': [initial_entry(MIXED)],
                'settings': {
                    'use_correct_form': True,
                    'scope': 'new',
                },
            },
        ]

        for scenario in scenarios:
            with self.subTest(settings=scenario['settings']):
                session = self.client.session
                session['doc_abbs'] = [mixed_card(), canonical_card()]
                session['initial_abbs'] = scenario['initial']
                session.save()

                payload = json.dumps(scenario['settings'])
                preview = self.client.post(
                    reverse('preview_abbreviation_table'),
                    data=payload,
                    content_type='application/json',
                )
                export = self.client.post(
                    reverse('make_abbreviation_table'),
                    data=payload,
                    content_type='application/json',
                )

                self.assertEqual(preview.status_code, 200)
                self.assertEqual(export.status_code, 200)
                self.assertEqual(
                    export['Content-Type'],
                    'application/vnd.openxmlformats-officedocument.'
                    'wordprocessingml.document',
                )

                preview_entries = preview.json()['entries']
                preview_names = [
                    entry['abbreviation'] for entry in preview_entries
                ]
                document = Document(io.BytesIO(export.content))
                export_names = [
                    row.cells[0].text
                    for row in document.tables[0].rows[1:]
                ]
                self.assertEqual(export_names, preview_names)

                if not scenario['settings']['use_correct_form']:
                    mixed = next(
                        entry for entry in preview_entries
                        if entry['abbreviation'] == MIXED
                    )
                    self.assertTrue(any(
                        part['mismatch']
                        for part in mixed['highlighted']
                    ))
