import io
import re
from dataclasses import dataclass
from typing import Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.shared import Cm, Pt, RGBColor
from docx.table import _Cell, Table

from .abbreviations import (
    Abbreviation,
    HighlightedCharacter,
    TableEntry,
    format_description,
    load_approved_dictionary,
)
from .extraction import (
    CharacterValidator,
    ExcludedSection,
    TextProcessor,
    process_abbreviations,
)


SECTION_PATTERNS = [
    'ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОПРЕДЕЛЕНИЯ ТЕРМИНОВ', 'СПИСОК СОКРАЩЕНИЙ'
]


class AbbreviationTableExtractor:
    """Class for extracting abbreviation table from a Word document."""
    NS = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    def __init__(self, section_patterns: List[str] = SECTION_PATTERNS):
        self.section_patterns = section_patterns

    def get_abbreviation_table(self, doc: Document) -> List[Abbreviation]:
        """Extract abbreviations table from document"""
        table_element = self._extract_table_element(doc)
        if table_element is None:
            return []

        return self._parse_table(table_element)

    def _extract_table_element(self, doc: Document) -> Optional[CT_Tbl]:
        """
        Extract the first table following a section matching `section_patterns`.
        Returns the block containing the table, or None if not found.
        """
        found_section = False

        for block in doc.element.body:
            if block.tag.endswith('p'): # paragraph block
                para_text = ''.join(
                    node.text for node in block.findall(
                        './/w:t', namespaces=self.NS
                    ) if node.text
                ).strip()

                if any(pattern.casefold() in para_text.casefold()
                    for pattern in self.section_patterns):

                    # Must NOT have a hyperlink (avoid ToC lines)
                    if block.find('.//w:hyperlink', namespaces=self.NS) is not None:
                        continue

                    # Must NOT end with a digit (avoid missformated ToC)
                    if para_text.strip().endswith(tuple("0123456789")):
                        continue

                    # Must have some heading indication (pStyle or outlineLvl)
                    para_style = block.find('.//w:pStyle', namespaces=self.NS)
                    outline_level = block.find('.//w:outlineLvl', namespaces=self.NS)
                    if para_style is not None or outline_level is not None:
                        found_section = True
                        continue

            if found_section and block.tag.endswith('tbl'):
                return block

        return None

    def _parse_table(self, table_element: CT_Tbl) -> List[Abbreviation]:
        """Parse table into list of abbreviation entries"""
        abb_entries: Dict[str, List[str]] = {}
        rows = table_element.findall('.//w:tr', namespaces=self.NS)

        # Get abbreviations and descriptions
        for idx, row in enumerate(rows):
            cell_values = [
                ''.join(
                    t.text for t in cell.findall('.//w:t', namespaces=self.NS) if t.text
                ).strip()
                for cell in row.findall('.//w:tc', namespaces=self.NS)
            ]
            if idx == 0 and set(cell_values[:2]) == {"Аббревиатура", "Расшифровка"}:
                continue

            if len(cell_values) == 2:
                abb, description = cell_values
                if abb in abb_entries:
                    if description not in abb_entries[abb]:
                        abb_entries[abb].append(description)
                else:
                    abb_entries[abb] = [description]

        return [
            {
                'abbreviation': abb,
                'descriptions': descriptions
            }
            for abb, descriptions in abb_entries.items()
        ]


class AbbreviationFormatter:
    """Class for formatting and sorting abbreviation table entries."""

    def format_table_entries(
        self,
        entries: List[TableEntry],
    ) -> List[TableEntry]:
        """Format descriptions without changing table row identity."""
        return [
            {
                'abbreviation': entry['abbreviation'].strip(),
                'description': format_description(
                    entry['abbreviation'].strip(),
                    entry['description'],
                ),
                'highlighted': entry.get('highlighted'),
            }
            for entry in entries
        ]

    def sort_abbreviations(
        self,
        abbreviations: List[TableEntry],
        sort_mode: str,
        script_order: str,
    ) -> List[TableEntry]:
        """Sort table entries by script group and configured row order."""
        def script_rank(abbreviation: str) -> int:
            for char in abbreviation:
                if re.match(r'[А-Яа-яЁё]', char):
                    return 0 if script_order == 'cyrillic_first' else 1
                if re.match(r'[A-Za-z]', char):
                    return 0 if script_order == 'latin_first' else 1
            return 2

        if sort_mode == 'appearance':
            return sorted(
                abbreviations,
                key=lambda entry: script_rank(entry['abbreviation']),
            )

        return sorted(
            abbreviations,
            key=lambda entry: (
                script_rank(entry['abbreviation']),
                entry['abbreviation'].casefold(),
                entry['description'].casefold(),
            ),
        )


class AbbreviationTableGenerator:
    """
    Class for generating formatted Word document tables with abbreviations.
    """

    def __init__(self):
        self.margins = {
            'top': 2.0,    # cm
            'bottom': 2.0, # cm
            'left': 3.0,   # cm
            'right': 1.5   # cm
        }
        self.first_column_width = 3.7  # cm
        self.font_name = 'Times New Roman'
        self.font_size = 12  # pt

    def generate_document(self, table_entries: List[TableEntry]) -> Document:
        """
        Generate a Word document with formatted abbreviation table.
        """
        doc = Document()

        # Set page margins
        for section in doc.sections:
            section.top_margin = Cm(self.margins['top'])
            section.bottom_margin = Cm(self.margins['bottom'])
            section.left_margin = Cm(self.margins['left'])
            section.right_margin = Cm(self.margins['right'])

            # Calculate second column width
            total_width = (
                section.page_width
                - section.left_margin
                - section.right_margin
            )
            self.second_column_width = total_width - Cm(self.first_column_width)

        # Create and format table
        table = self._create_table(doc, table_entries)
        self._set_column_widths(table)

        return doc

    def _create_table(
        self,
        doc: Document,
        table_entries: List[TableEntry],
    ) -> Table:
        """Create and format table with header and data rows."""
        # Create table with header
        table = doc.add_table(rows=1, cols=2)
        header_cells = table.rows[0].cells

        # Set header text
        header_cells[0].text = 'Аббревиатура'
        header_cells[1].text = 'Расшифровка'

        # Format header cells
        for cell in header_cells:
            self._format_cell(cell, bold=True)

        # Add and format data rows
        for entry in table_entries:
            row_cells = table.add_row().cells
            row_cells[0].text = entry['abbreviation']
            row_cells[1].text = entry['description']

            for cell in row_cells:
                self._format_cell(cell, bold=False)

        return table

    def _set_cell_border(self, cell: _Cell) -> None:
        """Set black borders on all four sides of the given cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        for edge in ("top", "bottom", "left", "right"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn('w:val'), 'single')    # single line
            element.set(qn('w:sz'), '4')          # line size
            element.set(qn('w:color'), '000000')  # black color
            tcPr.append(element)

    def _format_paragraph_spacing(self, cell: _Cell) -> None:
        """Set line spacing to single and remove spacing before/after paragraphs."""
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = Pt(self.font_size)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

    def _format_cell_text(self, cell: _Cell, bold: bool = False) -> None:
        """Format text in the given cell with specified font and style."""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
                run.font.bold = bold
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def _format_cell(self, cell: _Cell, bold: bool = False) -> None:
        """Apply all formatting to a cell."""
        self._set_cell_border(cell)
        self._format_paragraph_spacing(cell)
        self._format_cell_text(cell, bold)

    def _set_column_widths(self, table: Table) -> None:
        """Set fixed widths for table columns."""
        for row in table.rows:
            row.cells[0].width = Cm(self.first_column_width)
            row.cells[1].width = self.second_column_width



extractor = AbbreviationTableExtractor()
formatter = AbbreviationFormatter()
generator = AbbreviationTableGenerator()


@dataclass(frozen=True)
class ProcessedDocument:
    abbreviations: List[Abbreviation]
    initial_abbreviations: List[Abbreviation]


@dataclass(frozen=True)
class DisplayAbbreviation:
    source_abbreviation: str
    correct_form: Optional[str]
    display_abbreviation: str
    description: str
    highlighted: Optional[List[HighlightedCharacter]]


def find_bibliography_sections(file_path: str) -> List[ExcludedSection]:
    document = Document(file_path)
    return TextProcessor().extract_relevant_text(document).excluded_sections


def process_document(
    file_path: str,
    included_bibliography_sections: Optional[set[str]] = None,
) -> ProcessedDocument:
    dictionary = load_approved_dictionary()

    document = Document(file_path)
    initial_abbreviations = extractor.get_abbreviation_table(document)
    validator = CharacterValidator()
    for entry in initial_abbreviations:
        validation = validator.validate_abbreviation(
            entry['abbreviation'],
            dictionary,
        )
        if validation.get('correct_form'):
            entry['highlighted'] = validation.get('highlighted')

    text_processor = TextProcessor()
    text = text_processor.extract_relevant_text(
        document,
        included_section_ids=included_bibliography_sections,
    ).text
    abbreviations = process_abbreviations(text, dictionary)

    return ProcessedDocument(
        abbreviations=abbreviations,
        initial_abbreviations=initial_abbreviations,
    )


def prepare_abbreviation_table_entries(
    doc_abbs: List[Abbreviation],
    initial_abbreviations: Optional[List[Abbreviation]] = None,
    sort_mode: str = 'alphabetical',
    script_order: str = 'latin_first',
    use_correct_form: bool = True,
    scope: str = 'all',
) -> List[TableEntry]:
    """Build display rows shared by preview, comparison, and Word export."""
    selected_by_display: Dict[str, DisplayAbbreviation] = {}

    for entry in doc_abbs:
        description = entry.get('selected_description')
        if description is None:
            continue

        source_abbreviation = entry['abbreviation'].strip()
        correct_form = entry.get('correct_form')
        display_abbreviation = (
            correct_form
            if use_correct_form and correct_form
            else source_abbreviation
        )
        highlighted = (
            entry.get('highlighted')
            if display_abbreviation == source_abbreviation
            else None
        )
        candidate = DisplayAbbreviation(
            source_abbreviation=source_abbreviation,
            correct_form=correct_form,
            display_abbreviation=display_abbreviation,
            description=description,
            highlighted=highlighted,
        )

        existing = selected_by_display.get(display_abbreviation)
        if existing is None:
            selected_by_display[display_abbreviation] = candidate
            continue

        candidate_rank = (
            candidate.correct_form == candidate.display_abbreviation
            and candidate.source_abbreviation != candidate.display_abbreviation,
            candidate.source_abbreviation.casefold(),
            candidate.source_abbreviation,
            candidate.description.casefold(),
            candidate.description,
        )
        existing_rank = (
            existing.correct_form == existing.display_abbreviation
            and existing.source_abbreviation != existing.display_abbreviation,
            existing.source_abbreviation.casefold(),
            existing.source_abbreviation,
            existing.description.casefold(),
            existing.description,
        )
        if candidate_rank < existing_rank:
            selected_by_display[display_abbreviation] = candidate

    rows: List[TableEntry] = [
        {
            'abbreviation': entry.display_abbreviation,
            'description': entry.description,
            'highlighted': entry.highlighted,
        }
        for entry in selected_by_display.values()
    ]

    if scope == 'new':
        initial_names = {
            entry['abbreviation']
            for entry in (initial_abbreviations or [])
        }
        rows = [
            entry for entry in rows
            if entry['abbreviation'] not in initial_names
        ]

    formatted = formatter.format_table_entries(rows)
    return formatter.sort_abbreviations(
        formatted,
        sort_mode=sort_mode,
        script_order=script_order,
    )


def build_abbreviation_table_docx(
    abbreviations: List[TableEntry],
) -> bytes:
    document = generator.generate_document(abbreviations)

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()
