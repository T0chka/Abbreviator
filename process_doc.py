import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from thefuzz import process, fuzz

# Local imports
from extract_abbs import (
    get_all_abbreviations,
    compare_abbreviations
)
from format_abbs import clean_and_sort_abbreviations

ABB_DICT_PATH = 'data/abb_dict.csv'
DOC_PATH = "C:/Workspace/R-pharm/work/abbreviation_app/data/docs_examples"

# -----------------------------------------------------------------------------
# Abbreviations and context extraction
# -----------------------------------------------------------------------------

def load_abbreviation_dict():
    """
    Load abbreviations from CSV, or return empty DataFrame if not found.
    """
    if os.path.exists(ABB_DICT_PATH):
        return pd.read_csv(ABB_DICT_PATH)
    return pd.DataFrame(columns=["abbreviation", "description"])

def extract_relevant_text(doc):
    """
    Extracts text from the document, excluding sections like "СПИСОК ЛИТЕРАТУРЫ".
    The exclusion starts at the section title in bold or heading style
    and resumes at the next bold or heading section.
    """
    skip_sections = [
        "СПИСОК ЛИТЕРАТУРЫ",
        "Список использованной литературы",
        "Список использованных источников"
    ]

    relevant_text = []
    skip = False

    for para in doc.paragraphs:
        para_text = para.text.strip()

        is_heading = (
            para.style.name.startswith('Heading')
            or 'Заголовок' in para.style.name
        )
        is_bold = any(
            run.bold for run in para.runs if run.text.strip()
        )

        if any(t.upper() in para_text.upper() for t in skip_sections):
            print(
                f"[DEBUG] Detected: {para_text} - "
                f"Style: {para.style.name} - "
                f"Is Bold: {is_bold} - Is Heading: {is_heading}"
            )
        if ((is_bold or is_heading)
            and any(t.upper() in para_text.upper() for t in skip_sections)
            ):
            skip = True

        elif (is_bold or is_heading) and skip:
            print(
                f"[DEBUG] Resuming search at section: {para_text}"
            )
            skip = False

        if not skip:
            relevant_text.append(para_text)

    return " ".join(relevant_text)

def extract_abbs_from_text(text):
    """
    Extracts uppercase and mixed-case abbreviations from the document.
    Stops searching at "Список литературы" if the style is Heading 1.
    Excludes pure Roman numerals, specific terms, and words in quotes,
    and abbs that are 9 or more characters long and contain only letters.
    """
    doc_abbs = set()

    exclude_terms = {
        'ДИЗАЙН', 'ГЛАВНЫЙ', 'СПИСОК', 'ПРЯМОЙ', 'ПРИЕМ', 'ПРОТОКОЛ', 'ОТБОР',
        'КАЧЕСТВА', 'ПЕРИОД', 'ВЕДЕНИЕ', 'ЭТАП', 'ЭТИКА', 'СИНОПСИС', 'ЛИСТ',
        'ЦЕЛИ', 'РАБОТА', 'ИСТОРИЯ', 'ОЦЕНКА', 'СПОНСОР', 'ЗАДАЧИ', 'ДОСТУП',
        'КОНТРОЛЬ', 'ТЕРМИНОВ', 'ЗАПИСЕЙ', 'ГИПОТЕЗА', 'ДАННЫМИ'
    }
    roman_pattern = re.compile(
        r'^(?:[IVXLCDM]+(?:-[IVXLCDM]+)?)[A-Za-zА-Яа-яёЁ]*$',
        re.IGNORECASE
    )
    # Remove quoted words
    text_no_quotes = re.compile(r'«\S+»').sub('', text)
    words = text_no_quotes.split()

    # Find words with at least 2 uppercase (Latin or Cyrillic) letters
    matches = [
        word for word in words
        if re.search(r'[A-ZА-ЯЁ].*[A-ZА-ЯЁ]', word)
    ]
    
    for match in matches:
        clean_match = match.strip(':;,.»«][')

        # Remove '(' if at the beginning and ')' if unmatched
        if clean_match.startswith('('):
            clean_match = clean_match[1:]
        if clean_match.endswith(')') and clean_match.count('(') == 0:
            clean_match = re.sub(r'\)+$', '', clean_match)
        
        clean_match = clean_match.strip('»«][')

        # Exclude pure Roman numerals and specific excluded terms,
        # and overly long abbreviations
        if (not roman_pattern.match(clean_match)
            and clean_match not in exclude_terms
            and not (len(clean_match) > 8 and clean_match.isalpha())
            ):
            doc_abbs.add(clean_match)
    return doc_abbs

def find_abbreviation_context(doc, abbreviation, window=50, find_all=False):
    """
    Finds and returns snippets of text around occurrences of the abbreviation.
    If `find_all` is False, returns the first occurrence; True returns all.
    """
    contexts = set()    
    for para in doc.paragraphs:
        matches = re.finditer(
            rf'(?<!\w){re.escape(abbreviation)}(?!\w)', para.text
        )
        
        for match in matches:
            start = max(0, match.start() - window)
            end = min(len(para.text), match.end() + window)
            snippet = para.text[start:end]
            
            if find_all:
                contexts.add(snippet)
            else:
                return snippet
    
    return list(contexts) if find_all else None

# -----------------------------------------------------------------------------
# Mistyped Abbreviation Checking
# -----------------------------------------------------------------------------

def generate_all_mixed_forms(abb, cyr2lat, lat2cyr):
    """
    Given an abbreviation that may have mixed Cyrillic/Latin chars,
    generate all possible "conversion" permutations by:
      - leaving the character as is,
      - converting Cyrillic -> Latin,
      - converting Latin -> Cyrillic.
    """
    results = set()

    def backtrack(i, current):
        if i == len(abb):
            results.add("".join(current))
            return

        ch = abb[i]
        # 1) Leave the character as is
        current.append(ch)
        backtrack(i + 1, current)
        current.pop()

        # 2) Cyrillic -> Latin (if ch is in cyr2lat)
        if ch in cyr2lat:
            current.append(cyr2lat[ch])
            backtrack(i + 1, current)
            current.pop()

        # 3) Latin -> Cyrillic (if ch is in lat2cyr)
        if ch in lat2cyr:
            current.append(lat2cyr[ch])
            backtrack(i + 1, current)
            current.pop()

    backtrack(0, [])
    return results

def highlight_mixed_characters(abb, cyr2lat, lat2cyr):
    """
    Returns a debug-friendly string, highlighting each character from the dicts
    cyr2lat and lat2cyr with (Cyr) or (Lat).
    """
    highlighted_chars = []
    for ch in abb:
        if ch in cyr2lat:
            highlighted_chars.append(f"{ch}(Cyr)")
        elif ch in lat2cyr:
            highlighted_chars.append(f"{ch}(Lat)")
        else:
            highlighted_chars.append(ch)
    return "".join(highlighted_chars)

def check_and_correct_misstyped_abb(abb, abb_dict):
    """
    Detect if `abb` is a mixed-type (Cyr/Lat). If so, suggest
    one or more possible corrections from existing `abb_dict`.
    Returns the corrected abbreviation or the original if user skips.
    """
    # Map for character-by-character conversion
    cyr2lat = {
        'А': 'A',  # Cyrillic А -> Latin A
        'В': 'B',  # Cyrillic В -> Latin B
        'С': 'C',  # Cyrillic С -> Latin C
        'Е': 'E',  # Cyrillic Е -> Latin E
        'Н': 'H',  # Cyrillic Н -> Latin H
        'К': 'K',  # Cyrillic К -> Latin K
        'М': 'M',  # Cyrillic М -> Latin M
        'О': 'O',  # Cyrillic О -> Latin O
        'Р': 'P',  # Cyrillic Р -> Latin P
        'Т': 'T',  # Cyrillic Т -> Latin T
        'У': 'Y',  # Cyrillic У -> Latin Y
        'Х': 'X'   # Cyrillic Х -> Latin X
    }
    lat2cyr = {v: k for k, v in cyr2lat.items()}

    is_mixed = any(char in cyr2lat for char in abb) and \
               any(char in lat2cyr for char in abb)
    possible_forms = generate_all_mixed_forms(abb, cyr2lat, lat2cyr)
    abb_hl = highlight_mixed_characters(abb, cyr2lat, lat2cyr)

    matches = []

    # Check which generated forms already exist in abb_dict
    for form in possible_forms:
        row = abb_dict[abb_dict['abbreviation'] == form]
        if not row.empty:
            desc = row['description'].values[0]
            matches.append((form, desc))

    # No matches found; just return original
    if not matches and is_mixed:
        print(
            f"\n[INFO] No corrections found for mixed-language abbreviation: "
            f"'{abb_hl}'"
        )
        print(
            "[INFO] This may be a new abbreviation or mistyped. "
            "Consider adding it manually."
        )
        return abb
    
    # De-duplicate the matches
    matches = list(dict.fromkeys(matches))

    # If exactly 1 match, ask user to confirm
    if len(matches) == 1:
        corrected_abb, description = matches[0]
        corrected_abb_hl = highlight_mixed_characters(
            corrected_abb, cyr2lat, lat2cyr
        )
        print(f"\n[WARNING] '{abb_hl}' appears mistyped.")
        print(f"Did you mean '{corrected_abb_hl}'? - {description}")

        choice = input(
            f"Use '{corrected_abb}' instead of '{abb}'? (y/n): "
        ).strip().lower()
        return corrected_abb if choice == 'y' else abb

    # If multiple matches, prompt user to pick
    if matches:
        print(
            f"\n[WARNING] '{abb_hl}' appears mistyped. Possible corrections:"
        )
        for i, (corr, desc) in enumerate(matches, start=1):
            print(f" {i}. {corr} - {desc}")

        choice = input(
            f"Select the correct abbreviation for '{abb}' (1-{len(matches)}) "
            f"or Enter to skip: "
        ).strip()
        if choice.isdigit():
            if 1 <= int(choice) <= len(matches):
                corrected_abb, _ = matches[int(choice) - 1]
                confirm = input(
                    f"Use '{corrected_abb}' instead of '{abb}'? (y/n): "
                ).strip().lower()
                if confirm == 'y':
                    return corrected_abb
    return abb

# -----------------------------------------------------------------------------
# Custom description input
# -----------------------------------------------------------------------------

def get_similar_description(custom_desc, abb_dict, threshold=70):
    """
    Returns abbreviation-description pairs that have a similar description
    to the provided custom_desc, using fuzzy matching.
    """
    def normalize_text(text):
        """Normalize text by stripping, lowering, and keeping only letters."""
        text = text.strip().lower()
        text = re.sub(r'[^a-zа-яё\s]', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    normalized_custom_desc = normalize_text(custom_desc)
    abb_dict['norm_desc'] = abb_dict['description'].apply(normalize_text)

    matches = process.extract(
        normalized_custom_desc,
        abb_dict['norm_desc'],
        scorer=fuzz.token_sort_ratio,
        limit=5
    )
    abb_dict.drop(columns=['norm_desc'], inplace=True)
    
    # Filter by similarity threshold
    similar_pairs = [
        abb_dict.iloc[idx].to_dict()
        for _, score, idx in matches if score >= threshold
    ]
    return similar_pairs

def get_custom_description(abb, abb_dict, matched_abbs):
    """
    Prompt user to enter a custom description for an abbreviation.
    Check for similar descriptions and confirm with the user.
    Updates abb_dict and matched_abbs with the new description.
    """
    custom_desc = input(f"Enter custom description for '{abb}': ").strip()
    
    if custom_desc:
        # Check if the description already exists for another abbreviation
        similar_pairs = get_similar_description(custom_desc, abb_dict)
   
        if similar_pairs:
            print(f"\n[INFO] Similar descriptions found for '{abb}':")
            for i, pair in enumerate(similar_pairs, 1):
                print(f" {i}. {pair['abbreviation']} - {pair['description']}")

            choice = input(
                f"Use one of the existing pairs (1-{len(similar_pairs)}) "
                f"or enter 'n' to proceed with the new one: "
            ).strip().lower()

            if choice.isdigit() and 1 <= int(choice) <= len(similar_pairs):
                pair = similar_pairs[int(choice) - 1]
                print(
                    f"[INFO] Instead of '{abb}' using existing pair: "
                    f"{pair['abbreviation']} - {pair['description']}."
                )
                matched_abbs.drop(
                    matched_abbs[matched_abbs['abbreviation'] == abb].index,
                    inplace=True
                )
                matched_abbs = pd.concat(
                    [matched_abbs, pd.DataFrame([pair])],
                    ignore_index=True
                )
                # track_abb_dict("within get_custom_description, if selected similar")
                return matched_abbs, abb_dict

        # Add a new entry if no similar pair exists or user opts to continue
        new_entry = pd.DataFrame(
            [{'abbreviation': abb, 'description': custom_desc}]
        )
        abb_dict.loc[len(abb_dict)] = new_entry.iloc[0]

        matched_abbs.drop(
            matched_abbs[matched_abbs['abbreviation'] == abb].index,
            inplace=True
        )
        matched_abbs = pd.concat([matched_abbs, new_entry], ignore_index=True)
        print(
            f"[INFO] New pair '{abb} - {custom_desc}' "
            "added to abbreviation dictionary."
        )
        # track_abb_dict("within get_custom_description, after new pair added")
    else:
        print(f"[INFO] No description provided. '{abb}' will be excluded.")

    clean_and_sort_abbreviations(abb_dict)
    check_for_invalid_characters(abb_dict, stage="in get_custom_description")          
    return matched_abbs, abb_dict

# -----------------------------------------------------------------------------
# Handling Multiple Descriptions
# -----------------------------------------------------------------------------

def handle_multiple_descriptions(abb, descriptions, abb_dict, matched_abbs, doc):
    """
    Let user pick or create a new description if multiple exist.
    Uses get_custom_description() to update abb_dict with new entries.
    """
    print(f"\nMultiple descriptions found for '{abb}':")
    for i, desc in enumerate(descriptions, 1):
        print(f"{i}. {desc}")

    contexts = find_abbreviation_context(doc, abb, find_all=True)
    for context in contexts:
        print(f"   Context: ...{context}...")        

    print("0. Enter a custom description")
    while True:
        choice = input(
            f"Select the correct description for '{abb}' (1-{len(descriptions)})"
            " or 0 to add custom: "
        )
        if choice == "0":
            matched_abbs, abb_dict = get_custom_description(
                abb, abb_dict, matched_abbs
            )
            break
        
        elif choice.isdigit() and 1 <= int(choice) <= len(descriptions):
            chosen_desc = descriptions[int(choice) - 1]
            matched_abbs = matched_abbs[
                (matched_abbs['abbreviation'] != abb) |
                (matched_abbs['description'] == chosen_desc)
            ]
            break
        else:
            print("[WARNING] Invalid choice. Please try again.")

    return matched_abbs, abb_dict

# -----------------------------------------------------------------------------
# Core Logic: match and update abbreviations
# -----------------------------------------------------------------------------

def match_and_update_abbs(doc, doc_abbs, abb_dict):
    """
    Match abbreviations found in the doc with an existing dictionary, 
    prompting for new or custom descriptions as needed.
    Uses get_custom_description() to update abb_dict with new entries.
    """
    # 1) Match existing abbreviations in the dictionary
    matched_abbs = abb_dict[abb_dict['abbreviation'].isin(doc_abbs)].copy()
    
    new_abbs = doc_abbs - set(matched_abbs['abbreviation'])

    # Handle abbreviations with multiple descriptions
    for abb in matched_abbs['abbreviation'].unique():
        descriptions = abb_dict.loc[
            abb_dict['abbreviation'] == abb, 'description'
        ].unique()
    
        if len(descriptions) > 1:
            matched_abbs, abb_dict = handle_multiple_descriptions(
                abb, descriptions, abb_dict, matched_abbs, doc
            )
    
    # 2) Prompt for new abbreviations not in dictionary
    if new_abbs:
        for abb in new_abbs:

            # Check for possible mistypes
            corrected_abb = check_and_correct_misstyped_abb(abb, abb_dict)
            
            # If corrected_abb exists, skip new description prompt
            if corrected_abb in abb_dict['abbreviation'].values:
                rows_for_corrected = abb_dict[
                    abb_dict['abbreviation'] == corrected_abb
                ]
                matched_abbs = pd.concat(
                    [matched_abbs, rows_for_corrected], ignore_index=True
                )

                # If multiple descriptions exist for the corrected abbreviation
                descriptions = rows_for_corrected['description'].unique()    
                if len(descriptions) > 1:
                    matched_abbs, abb_dict = handle_multiple_descriptions(
                        corrected_abb, descriptions, abb_dict, matched_abbs, doc
                    )
                continue
            
            # Otherwise, prompt user for a description
            contexts = find_abbreviation_context(doc, abb, find_all=True)
            print(f"\nFound '{abb}' in text:")
            for context in contexts:
                print(f"...{context}...")
            matched_abbs, abb_dict = get_custom_description(
                abb, abb_dict, matched_abbs
            )

    clean_and_sort_abbreviations(matched_abbs)
    return matched_abbs, abb_dict

def search_abbs_in_text(doc, matched_abbs, abb_dict):
    """
    Search for abbreviations from abb_dict that may have been missed
    in the initial detection and add them to matched_abbs.
    """
    stop_section="Список литературы"
    found_text = []
    
    for para in doc.paragraphs:
        if (stop_section.lower() in para.text.lower() and
            para.style.name == 'Heading 1'):
            break
        found_text.append(para.text)

    found_text = " ".join(found_text)
    
    for abb in abb_dict['abbreviation']:
        if (re.search(rf'(?<!\w){re.escape(abb)}(?!\w)', found_text) and
            abb not in doc_abbs):
            matched_abbs = pd.concat(
                [matched_abbs, abb_dict[abb_dict['abbreviation'] == abb]],
                ignore_index=True
            )

    clean_and_sort_abbreviations(matched_abbs)
    return matched_abbs

def review_ole(matched_abbs, doc, abb_dict):
    """
    Review one-letter abbreviations (ola), prompting users to confirm 
    or exclude them from the table.
    """
    ola_abbs = matched_abbs[
        matched_abbs['abbreviation'].str.len() == 1
    ].copy()
    
    if not ola_abbs.empty:
        print("\n[INFO] One-letter abbreviations detected:")
        
        to_keep = []
        for abb in ola_abbs['abbreviation'].unique():
            descriptions = ola_abbs[ola_abbs['abbreviation'] == abb]
            
            if len(descriptions) > 1:
                print(f"\nAbbreviation: '{abb}' has multiple descriptions:")
                for desc in descriptions['description'].unique():
                    print(f"- {desc}")
                                
                # Ask if user wants to handle or skip
                choice = input(
                    f"Do you want to add '{abb}' to the table? (y/n): "
                ).strip().lower()
                
                if choice == 'y':
                    matched_abbs, abb_dict = handle_multiple_descriptions(
                        abb, descriptions['description'].unique(),
                        abb_dict, matched_abbs, doc
                    )
                    to_keep.append(abb)
                else:
                    print(f"[INFO] '{abb}' will be excluded.")
            else:
                desc = descriptions['description'].values[0]
                print(f"\nAbbreviation: '{abb}'")
                print(f" - {desc}")
                choice = input(
                    f"Add this description for '{abb}'? (y/n/c):"
                ).strip().lower()
                
                if choice == 'y':
                    to_keep.append(abb)
                elif choice == 'c':
                    matched_abbs, abb_dict = get_custom_description(
                        abb, abb_dict, matched_abbs
                    )
                    to_keep.append(abb)
                else:
                    print(f"[INFO] '{abb}' will be excluded.")

        matched_abbs = matched_abbs[
            (matched_abbs['abbreviation'].str.len() > 1) | 
            (matched_abbs['abbreviation'].isin(to_keep))
        ]
        clean_and_sort_abbreviations(matched_abbs)

    return matched_abbs, abb_dict

# -----------------------------------------------------------------------------
# Output generation
# -----------------------------------------------------------------------------

def set_cell_border(cell):
    """
    Sets black borders on all four sides of the given cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ("top", "bottom", "left", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn('w:val'), 'single')    # single line
        element.set(qn('w:sz'), '4')          # line size
        element.set(qn('w:color'), '000000')  # black color
        tcPr.append(element)

def format_paragraph_spacing(cell):
    """
    Sets line spacing to single and spacing after to 0 for paragraphs in a cell.
    """
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = Pt(12)  # Single line spacing
        paragraph.paragraph_format.space_after = Pt(0)    # No space after paragraph
        paragraph.paragraph_format.space_before = Pt(0)   # No space before paragraph

def format_cell_text(cell, bold=False):
    """
    Formats text in the given cell with specified font and style.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

def generate_abbreviation_table(matched_abbs):
    """
    Generates a .docx file with only a table,
    with page margins set to top/bottom=2cm, left=3cm, right=1.5cm,
    and a fixed 3.7cm width for the first column.
    """
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)

        total_width = (
            section.page_width - section.left_margin - section.right_margin
        )
        second_col_width = total_width - Cm(3.7)

    # Create table (header row + body, 2 columns)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Аббревиатура'
    hdr_cells[1].text = 'Расшифровка'

    # Format header
    for cell in hdr_cells:
        set_cell_border(cell)
        format_paragraph_spacing(cell)
        format_cell_text(cell, bold=True)

    # Table body
    for _, row_data in matched_abbs.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row_data['abbreviation']
        row_cells[1].text = str(row_data['description'])

        for cell in row_cells:
            set_cell_border(cell)
            format_paragraph_spacing(cell)
            format_cell_text(cell, bold=False)

    # Enforce column widths
    for row in table.rows:
        row.cells[0].width = Cm(3.7)
        row.cells[1].width = second_col_width

    # Save
    out_path = 'output/abbreviations_table.docx'
    os.makedirs('output', exist_ok=True)
    doc.save(out_path)
    print(f"Abbreviation table saved to {out_path}")

# -----------------------------------------------------------------------------
# Additional Checks
# -----------------------------------------------------------------------------

def check_inconsistencies(matched_abbs):
    """
    1) Print how many abbreviations have more than one unique description.
    2) Print how many descriptions are used by more than one abbreviation.
    """
    # Abbreviations with multiple unique descriptions
    multi_desc = (matched_abbs.groupby('abbreviation')['description']
                            .nunique()
                            .reset_index()
                            .query('description > 1'))
    count_mult_desc = len(multi_desc)
    if count_mult_desc > 0:
        print(
            f"\n[INFO] {count_mult_desc} abbreviation(s) have "
            "multiple descriptions:"
        )
        duplicates = matched_abbs[
            matched_abbs['abbreviation'].isin(multi_desc['abbreviation'])
        ]
        print(duplicates)
    else:
        print("\n[INFO] No abbreviations with multiple descriptions.")

    # Descriptions used by multiple abbreviations
    multi_abbs = (matched_abbs.groupby('description')['abbreviation']
                             .nunique()
                             .reset_index()
                             .query('abbreviation > 1'))
    count_mult_abbs = len(multi_abbs)
    if count_mult_abbs > 0:
        print(
            f"\n[INFO] {count_mult_abbs} description(s) are shared by "
            "multiple abbreviations:"
        )
        duplicates = matched_abbs[
            matched_abbs['description'].isin(multi_abbs['description'])
        ]
        print(duplicates)
    else:
        print("\n[INFO] No descriptions are shared by multiple abbreviations.")

def review_new_entries(abb_dict_before, abb_dict_after):
    """
    Identify newly added rows (by abbreviation+description) in abb_dict,
    and ask user whether to save.
    """
    merged = abb_dict_after.merge(
        abb_dict_before,
        on=['abbreviation', 'description'],
        how='left',
        indicator=True
    )
    new_entries = merged[merged['_merge'] == 'left_only'][
        ['abbreviation', 'description']
    ]

    if not new_entries.empty:
        print(
            "\n[INFO] The following new entries have been identified "
            "in the abbreviation dictionary:"
        )
        print(new_entries)

        confirm = input(
            "\nAre you sure you want to permanently update the dictionary "
            "with these new entries? (y/n): "
        ).strip().lower()
        if confirm == 'y':
            abb_dict_after.to_csv(
                ABB_DICT_PATH, index=False, encoding='utf-8-sig'
            )
            print("[INFO] Abbreviation dictionary updated.")
        else:
            print("[INFO] No changes will be written.")
    else:
        print("\n[INFO] No new entries found to update the dictionary.")

# -----------------------------------------------------------------------------
# Debugging
# -----------------------------------------------------------------------------

def track_abb_dict(label):
        print(f"\n[DEBUG] After {label}:")
        print(f" - Is abb_dict a view? {abb_dict._is_view}")
        print(f" - Is abb_dict a copy? {abb_dict._is_copy is not None}")
        print(f" - Shape: {abb_dict.shape}")
        print(abb_dict)

def check_for_invalid_characters(df, stage="Unknown"):
    invalid_rows = df[
        df['description'].str.contains(r'[\x00-\x1F\x7F]', na=False)
    ]
    if not invalid_rows.empty:
        print(f"[ERROR - {stage}] Invalid descriptions found:\n", invalid_rows)
        raise ValueError(
            f"Control characters detected after {stage}. Please clean the data."
        )

# -----------------------------------------------------------------------------
# Main script
# -----------------------------------------------------------------------------

if __name__ == "__main__":
    doc_file = [f for f in os.listdir(DOC_PATH)
                if f.endswith('.docx') and not f.startswith('~$')]
    if len(doc_file) != 1:
        raise FileNotFoundError("Expected exactly one .docx file in the folder.")
    
    doc_path = os.path.join(DOC_PATH, doc_file[0])
    print(f"[INFO] Processing document: {doc_path}")

    doc = Document(doc_path)
    text = extract_relevant_text(doc)
    doc_abbs = extract_abbs_from_text(text)
    
    abb_dict = load_abbreviation_dict()
    # track_abb_dict("just loaded")
    matched_abbs, abb_dict_updated = match_and_update_abbs(doc, doc_abbs, abb_dict)
    # track_abb_dict("after match_and_update_abbs")
    matched_abbs = search_abbs_in_text(doc, matched_abbs, abb_dict)
    # track_abb_dict("after search_abbs_in_text")
    matched_abbs, abb_dict_updated = review_ole(matched_abbs, doc, abb_dict_updated)
    # track_abb_dict("after review_ole")
    
    
    # Generate abbreviation table
    generate_abbreviation_table(matched_abbs)
    # track_abb_dict("after generate_abbreviation_table")

    # Compare changes
    initial_abbs = get_all_abbreviations(DOC_PATH)    
    compare_abbreviations(old_abbs=initial_abbs, new_abbs=matched_abbs)

    # Check for abbreviations with multiple unique descriptions
    check_inconsistencies(matched_abbs)

    # Identify new entries that were not in the original dictionary for user review
    review_new_entries(abb_dict, abb_dict_updated)